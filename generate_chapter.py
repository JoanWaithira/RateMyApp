"""
Generate a thesis-style Word chapter for the usability evaluation of the
Energy Digital Twin for Smart Building Management.
Run:  python generate_chapter.py
Output: Chapter_5_Usability_Evaluation.docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── helpers ────────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_fmt(para, first_indent=None, space_before=None,
             space_after=None, line_spacing=None, alignment=None):
    pf = para.paragraph_format
    if first_indent  is not None: pf.first_line_indent = Cm(first_indent)
    if space_before  is not None: pf.space_before       = Pt(space_before)
    if space_after   is not None: pf.space_after        = Pt(space_after)
    if line_spacing  is not None: pf.line_spacing        = Pt(line_spacing)
    if alignment     is not None: para.alignment         = alignment

def add_heading(doc, text, level=1):
    """Add a styled heading that matches academic thesis conventions."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(24)
        pf.space_after  = Pt(12)
        run = p.add_run(text)
        set_font(run, size=16, bold=True)
    elif level == 2:
        pf.space_before = Pt(18)
        pf.space_after  = Pt(8)
        run = p.add_run(text)
        set_font(run, size=14, bold=True)
    elif level == 3:
        pf.space_before = Pt(12)
        pf.space_after  = Pt(6)
        run = p.add_run(text)
        set_font(run, size=12, bold=True, italic=True)
    return p

def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    para_fmt(p,
             first_indent=1.25 if indent else 0,
             space_before=0, space_after=6, line_spacing=22)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_font(run)
    return p

def add_body_noindent(doc, text):
    return add_body(doc, text, indent=False)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    pf.line_spacing = Pt(18)
    run = p.add_run(text)
    set_font(run)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after  = Pt(2)
    pf.line_spacing = Pt(18)
    run = p.add_run(text)
    set_font(run)
    return p

def shade_row(row, hex_color="E8EFF7"):
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  hex_color)
        tcPr.append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = "Table Grid"
    # header row
    hdr = table.rows[0]
    shade_row(hdr, "2563EB")
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_font(run, size=11, bold=True, color=(255, 255, 255))
    # data rows
    for ri, row_data in enumerate(rows):
        row_obj = table.rows[ri + 1]
        if ri % 2 == 1:
            shade_row(row_obj, "F0F4FF")
        for ci, cell_text in enumerate(row_data):
            cell = row_obj.cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_font(run, size=10.5)
    # column widths
    if col_widths:
        for row_obj in table.rows:
            for i, w in enumerate(col_widths):
                row_obj.cells[i].width = Cm(w)
    doc.add_paragraph()  # spacing after table
    return table

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_fmt(p, space_before=2, space_after=10, line_spacing=16)
    run = p.add_run(text)
    set_font(run, size=10, italic=True)
    return p

def add_spacer(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        para_fmt(p, space_before=0, space_after=0, line_spacing=6)


# ── document setup ─────────────────────────────────────────────────────────────

doc = Document()

# page margins
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.5)   # wider left for binding
    section.right_margin  = Cm(2.5)

# ── CHAPTER TITLE ──────────────────────────────────────────────────────────────

add_spacer(doc, 2)
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_fmt(title_p, space_before=0, space_after=4)
t1 = title_p.add_run("CHAPTER 5")
set_font(t1, size=13, bold=True, color=(37, 99, 235))

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_fmt(sub_p, space_before=4, space_after=20)
t2 = sub_p.add_run(
    "Usability Evaluation of the Energy Digital Twin\n"
    "for Smart Building Management"
)
set_font(t2, size=20, bold=True)

divider = doc.add_paragraph()
divider.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_fmt(divider, space_before=0, space_after=24)
d_run = divider.add_run("─" * 60)
set_font(d_run, size=10, color=(148, 163, 184))

add_spacer(doc)

# ── 5.1  Introduction ──────────────────────────────────────────────────────────

add_heading(doc, "5.1  Introduction", level=1)

add_body(doc,
    "The development of the Energy Digital Twin (EDT) for the GATE Institute "
    "building represented a significant step toward intelligent, data-driven "
    "building management. However, a technological artefact — however "
    "sophisticated — derives its practical value only insofar as its intended "
    "users can interact with it effectively and derive actionable insight from "
    "it. Accordingly, this chapter presents the design, rationale, and "
    "execution of a structured usability evaluation conducted to assess whether "
    "the EDT platform meets the information needs of a diverse range of "
    "building stakeholders."
)
add_body(doc,
    "The evaluation was conducted by the researcher, Joan Waithira Njoroge, "
    "as part of the Master of Science in Spatial Engineering programme at the "
    "University of Twente. The study combined quantitative feature-level "
    "ratings with qualitative task observations and open-ended feedback, "
    "producing a multidimensional picture of user experience. An AI-assisted "
    "summarisation layer was incorporated into the survey instrument to provide "
    "immediate, participant-specific recaps that could support iterative "
    "reflection during data collection."
)
add_body(doc,
    "The remainder of this chapter is organised as follows. Section 5.2 "
    "situates the evaluation within the broader research objectives. "
    "Section 5.3 describes the study design and methodological choices. "
    "Section 5.4 details the survey instrument and its constituent stages. "
    "Section 5.5 characterises the participant sample. Section 5.6 presents "
    "the procedural protocol. Section 5.7 describes the feature evaluation "
    "framework. Section 5.8 outlines the data collection and storage "
    "infrastructure. Section 5.9 addresses ethical considerations, and "
    "Section 5.10 summarises the chapter."
)

# ── 5.2  Research Objectives ───────────────────────────────────────────────────

add_heading(doc, "5.2  Research Objectives and Evaluation Goals", level=1)

add_body(doc,
    "The overarching research question guiding this phase of the study is: "
    "To what extent does the Energy Digital Twin platform support the "
    "information needs of diverse building stakeholders in understanding, "
    "monitoring, and acting upon building energy performance data? This "
    "question was operationalised into four specific evaluation objectives:"
)

add_numbered(doc,
    "Assess the navigability and intuitiveness of the 3D building interface "
    "across stakeholder groups with varying levels of technical expertise."
)
add_numbered(doc,
    "Evaluate the clarity and communicative effectiveness of the platform's "
    "eight core feature modules — namely the 3D Building Viewer, Energy "
    "Monitoring, Solar & Battery, Room Air Quality (IAQ), Fault Detection, "
    "Scenarios, Forecast, and Role-Specific Dashboard."
)
add_numbered(doc,
    "Determine the perceived utility of AI-assisted decision support, "
    "specifically the conversational assistant embedded within the platform, "
    "in helping stakeholders interpret data and identify recommended actions."
)
add_numbered(doc,
    "Identify the features most and least valued by participants, and collect "
    "qualitative evidence of usability barriers that may inform future "
    "development iterations."
)

add_body(doc,
    "These objectives were designed to generate both summative evidence — "
    "suitable for validating the current design — and formative evidence that "
    "can guide iterative refinement of the platform prior to operational "
    "deployment."
)

# ── 5.3  Study Design ──────────────────────────────────────────────────────────

add_heading(doc, "5.3  Study Design", level=1)

add_heading(doc, "5.3.1  Research Approach", level=2)

add_body(doc,
    "A mixed-methods, task-based usability study was employed. This design "
    "combines structured task scenarios — in which participants are asked to "
    "complete representative activities within the platform — with Likert-scale "
    "feature ratings, forced-choice questions, and free-text open responses. "
    "This combination is consistent with established practice in user experience "
    "research (Nielsen, 1994; Rubin & Chisnell, 2008) and aligns with the "
    "ISO 9241-11:2018 definition of usability as the degree to which a product "
    "can be used by specified users to achieve specified goals with "
    "effectiveness, efficiency, and satisfaction."
)
add_body(doc,
    "A key design decision was to deliver the evaluation in an unmoderated, "
    "remote format. Participants accessed the survey instrument and the live "
    "EDT platform independently, without a researcher present. This choice "
    "was motivated by two considerations. First, it enabled participation by "
    "stakeholders distributed across organisational and geographic contexts, "
    "including facilities managers, sustainability officers, and external "
    "visitors who could not conveniently be co-located with the researcher. "
    "Second, an unmoderated setting reduces social desirability bias — "
    "participants are less likely to moderate their responses based on "
    "perceived researcher expectations when completing the survey privately."
)

add_heading(doc, "5.3.2  Relationship to the Broader Research Design", level=2)

add_body(doc,
    "This usability evaluation constitutes the validation phase of the "
    "research and sits downstream of the design and implementation phases "
    "described in earlier chapters. The evaluation is not intended to produce "
    "generalisable findings about digital twins in general; rather, it serves "
    "to validate a specific artefact — the GATE Institute EDT — against the "
    "functional and communicative requirements identified during the design "
    "phase. The sample size and participant selection were therefore guided by "
    "purposive criteria (representation of the platform's intended stakeholder "
    "groups) rather than probabilistic criteria."
)

# ── 5.4  The Survey Instrument ─────────────────────────────────────────────────

add_heading(doc, "5.4  The Survey Instrument", level=1)

add_body(doc,
    "A bespoke, browser-based survey application — RateMyApp — was developed "
    "specifically for this evaluation and hosted via Netlify at a publicly "
    "accessible URL. The instrument was structured into seven sequential stages, "
    "each serving a distinct methodological purpose. Table 5.1 provides an "
    "overview of the stages and their primary data collection function."
)

add_table(doc,
    headers=["Stage", "Title", "Primary Function"],
    rows=[
        ["1", "Participant Information & Consent",
         "Ethical disclosure; informed consent capture"],
        ["2", "Role & Affiliation Selection",
         "Participant stratification by stakeholder type and organisation"],
        ["3", "Guided Platform Tour",
         "Contextual orientation; exposure to all major features before tasks"],
        ["4", "Task-Based Testing",
         "Behavioural evidence of navigability across three guided tasks"],
        ["5", "Feature Ratings",
         "Quantitative usefulness ratings (1–5 Likert) per feature module"],
        ["6", "Open Feedback",
         "Qualitative evidence: impressions, barriers, adoption intent"],
        ["7", "Thank You & AI Summary",
         "Participant-facing recap; AI summary saved to research database"],
    ],
    col_widths=[1.5, 5.5, 8.5]
)
add_caption(doc, "Table 5.1. Structure of the RateMyApp survey instrument.")

add_heading(doc, "5.4.1  Stage 1: Participant Information and Informed Consent", level=2)

add_body(doc,
    "The first stage presented participants with a structured participant "
    "information sheet. In compliance with the ethical guidelines of the "
    "University of Twente and the principles of the Declaration of Helsinki, "
    "the information sheet detailed the purpose of the study, the nature of "
    "participation, the anonymisation of data, the voluntary nature of "
    "participation, and participants' right to withdraw at any stage without "
    "consequence. Consent was captured via an explicit checkbox, and only "
    "participants who actively confirmed their agreement could proceed to "
    "subsequent stages. No data beyond the consent confirmation itself was "
    "recorded prior to the checkbox being checked."
)

add_heading(doc, "5.4.2  Stage 2: Role and Affiliation Selection", level=2)

add_body(doc,
    "Participants were asked to self-identify their professional role from a "
    "predefined list of six stakeholder archetypes identified during the "
    "platform's requirements analysis phase. These archetypes were: Building "
    "Director or Management, Facilities Manager, Sustainability Officer, IT "
    "Staff, Office Worker, and External Visitor or Stakeholder. Role "
    "self-identification served two purposes: it enabled stratified analysis "
    "of responses by stakeholder group, and it informed the adaptive content "
    "in Stage 4, where task framing could be contextualised to the "
    "participant's declared role."
)
add_body(doc,
    "Participants additionally indicated their organisational affiliation — "
    "University of Twente, GATE Institute, or Other — providing a secondary "
    "stratification variable for analysis."
)

add_heading(doc, "5.4.3  Stage 3: Guided Platform Tour", level=2)

add_body(doc,
    "Prior to undertaking the structured tasks, participants were presented "
    "with a guided orientation stage. An embedded live preview of the EDT "
    "platform was made available via an inline iframe, with a direct link to "
    "open the platform in a full browser tab for participants who preferred a "
    "larger viewing environment. The guided tour stage encouraged participants "
    "to explore the platform for approximately two minutes, with explicit "
    "prompts to: (1) scan the main navigation and replay panel; (2) open at "
    "least two feature tabs; and (3) identify which view seemed most "
    "immediately relevant to their declared role."
)
add_body(doc,
    "The tour stage was methodologically important for two reasons. First, it "
    "reduced the cognitive load associated with encountering an unfamiliar "
    "interface cold before being asked to perform structured tasks. Second, it "
    "provided each participant with comparable baseline exposure to the "
    "platform's structure, mitigating the effect of prior familiarity "
    "differences on task performance ratings."
)

add_heading(doc, "5.4.4  Stage 4: Task-Based Testing", level=2)

add_body(doc,
    "Stage 4 constituted the primary behavioural component of the evaluation. "
    "Participants were guided through three structured tasks that represented "
    "core use cases of the EDT platform. Each task was presented with a "
    "step-by-step checklist of suggested actions, a follow-up question "
    "assessing ease of completion, and an optional open-text field for "
    "qualitative notes. Table 5.2 summarises the three tasks and their "
    "associated follow-up questions."
)

add_table(doc,
    headers=["Task", "Heading", "Follow-Up Question", "Response Options"],
    rows=[
        ["1",
         "Explore the Building and Select a View",
         "Could you find and navigate to different areas easily?",
         "Very easy / Somewhat easy / Difficult"],
        ["2",
         "Inspect Building Conditions Using Dashboards and Heatmaps",
         "Was it clear what the building conditions are?",
         "Very clear / Somewhat clear / Unclear"],
        ["3",
         "Use the AI Assistant for Decision Support",
         "Did the AI insights help you understand what actions might be needed?",
         "Very helpful / Somewhat helpful / Not helpful"],
    ],
    col_widths=[1.0, 4.5, 5.5, 4.5]
)
add_caption(doc, "Table 5.2. Task-based evaluation scenarios and follow-up questions.")

add_body(doc,
    "Task 1 assessed three-dimensional spatial navigation — the ability to "
    "orient oneself within the 3D building model and locate specific areas. "
    "This task directly evaluated one of the platform's foundational "
    "capabilities and is of particular relevance to stakeholders whose role "
    "involves physical building management."
)
add_body(doc,
    "Task 2 assessed the participant's ability to use the dashboard and "
    "heatmap panels to understand live and replayed building conditions, "
    "including energy consumption by circuit, and indoor air quality metrics "
    "(temperature, humidity, and CO₂) by room. This task evaluated the "
    "communicative clarity of the platform's data visualisations."
)
add_body(doc,
    "Task 3 assessed the perceived utility of the conversational AI assistant "
    "embedded within the platform. Participants were instructed to ask the "
    "assistant questions about the building — including spatially anchored "
    "queries (e.g., 'show me elevators') and data-oriented queries (e.g., "
    "'current temperature in the conference room') — and to note whether the "
    "AI's responses helped them identify actionable next steps. This task was "
    "designed to assess the value-add of AI-mediated decision support over "
    "direct dashboard inspection."
)

add_heading(doc, "5.4.5  Stage 5: Feature Ratings", level=2)

add_body(doc,
    "Having interacted with the platform through the tour and structured tasks, "
    "participants were asked to rate each of the eight core feature modules on "
    "a five-point Likert scale anchored at 1 (not useful) and 5 (very useful). "
    "An optional free-text comment field was provided for each feature, "
    "allowing participants to specify what worked well, what was unclear, or "
    "what should be changed. Participants additionally completed two dropdown "
    "selections: which feature they would use most often in their daily work, "
    "and which feature they considered in most need of improvement."
)
add_body(doc,
    "The eight features evaluated were: (1) 3D Building Viewer, (2) Energy "
    "Monitoring, (3) Solar & Battery, (4) Room Air Quality, (5) Fault "
    "Detection, (6) Scenarios, (7) Forecast, and (8) Role-Specific Dashboard. "
    "These features correspond directly to the functional modules implemented "
    "in the platform and to the user stories collected during the requirements "
    "analysis phase of the research."
)

add_heading(doc, "5.4.6  Stage 6: Open Feedback", level=2)

add_body(doc,
    "The penultimate stage collected holistic impressions of the platform "
    "through a combination of structured and open-ended questions. An overall "
    "satisfaction rating was captured using a five-point emoji-anchored scale, "
    "labelled from 'Needs a lot of work' to 'Exceptional'. Participants were "
    "then asked four open-text questions: what they found most useful or "
    "impressive; what they found confusing, missing, or frustrating; whether "
    "they would use the system in their daily work (four-option forced choice); "
    "and any additional comments, ideas, or questions."
)
add_body(doc,
    "Participants were additionally offered the option to provide their name "
    "and email address. These fields were explicitly marked as optional and "
    "were included solely to facilitate follow-up contact if participants "
    "wished to share further thoughts with the researcher. The consent "
    "information presented in Stage 1 explicitly stated that no identifying "
    "information was required for participation."
)

add_heading(doc, "5.4.7  Stage 7: AI-Generated Summary and Database Storage", level=2)

add_body(doc,
    "Upon survey submission, participants were presented with a personalised "
    "AI-generated summary of their responses on the thank-you page. This "
    "summary was produced by the OpenAI language model via the application's "
    "backend (a Netlify serverless function), drawing on the participant's "
    "role, task outcomes, feature ratings, and open-text comments to construct "
    "a concise, research-style recap. In cases where the AI generation endpoint "
    "was unavailable, a deterministic fallback function generated a structured "
    "narrative from the same data fields."
)
add_body(doc,
    "All responses — including the AI-generated summary — were stored in a "
    "Supabase PostgreSQL database via a secure serverless function. The "
    "database used row-level security policies and service-role authentication "
    "to prevent unauthorised direct access. A dual-database architecture was "
    "implemented, with a primary and a secondary Supabase instance, to provide "
    "data redundancy. The survey responses were stored as structured records "
    "containing task outcomes, individual feature ratings, open-text responses, "
    "timestamps, and AI summaries."
)

# ── 5.5  Participants ──────────────────────────────────────────────────────────

add_heading(doc, "5.5  Participant Sample", level=1)

add_heading(doc, "5.5.1  Sampling Strategy", level=2)

add_body(doc,
    "Participants were recruited through purposive sampling, targeting "
    "individuals with professional or academic connections to the GATE "
    "Institute building and to smart building management more broadly. "
    "Recruitment channels included direct outreach within the University of "
    "Twente, communication with GATE Institute staff, and distribution to "
    "professional contacts in the facilities management and sustainability "
    "sectors. The survey was designed to be accessible to participants "
    "without prior familiarity with the EDT platform, ensuring that usability "
    "findings would reflect first-exposure experiences rather than trained "
    "user performance."
)

add_heading(doc, "5.5.2  Stakeholder Roles and Affiliations", level=2)

add_body(doc,
    "The six stakeholder archetypes included in the survey were informed by "
    "the user role analysis conducted during the platform design phase. Each "
    "archetype represents a distinct set of information needs and interaction "
    "patterns with building management data. Table 5.3 summarises the "
    "archetypes and their characteristic perspectives."
)

add_table(doc,
    headers=["Role", "Primary Perspective", "Key Information Needs"],
    rows=[
        ["Building Director / Management",
         "Strategic decisions, costs, sustainability",
         "Energy cost, carbon performance, capital investment cases"],
        ["Facilities Manager",
         "Day-to-day operations, maintenance, equipment",
         "Fault alerts, circuit-level anomalies, equipment status"],
        ["Sustainability Officer",
         "Carbon reporting, energy targets, green goals",
         "Solar yield, consumption trends, scenario modelling"],
        ["IT Staff",
         "Systems, servers, infrastructure",
         "Data connectivity, system reliability, sensor status"],
        ["Office Worker",
         "Daily comfort, indoor environment",
         "Room temperature, CO₂ levels, humidity conditions"],
        ["External Visitor / Stakeholder",
         "Project review, external oversight",
         "Overall platform legibility, headline metrics"],
    ],
    col_widths=[4.5, 5.0, 6.0]
)
add_caption(doc, "Table 5.3. Stakeholder archetypes and characteristic information needs.")

add_body(doc,
    "This role diversity was an intentional design choice. An energy digital "
    "twin intended for a multi-stakeholder building environment must serve users "
    "whose technical literacy, data fluency, and task priorities differ "
    "substantially. A usability evaluation that sampled from a single "
    "stakeholder group — for example, only facilities managers — would "
    "generate findings of limited generalisability to the platform's intended "
    "user population."
)

# ── 5.6  Procedure ─────────────────────────────────────────────────────────────

add_heading(doc, "5.6  Evaluation Procedure", level=1)

add_heading(doc, "5.6.1  Access and Orientation", level=2)

add_body(doc,
    "Participants received a short briefing message — distributed via email "
    "or direct communication — explaining that they were invited to evaluate "
    "a digital building management platform as part of a Master's thesis "
    "research project. The message provided a direct link to the survey URL "
    "and confirmed that the evaluation was expected to take approximately "
    "five to ten minutes to complete, with an additional two to five minutes "
    "of platform exploration time."
)
add_body(doc,
    "No preparatory materials, tutorials, or prior exposure to the platform "
    "were provided before the survey. This design decision was deliberate: "
    "the evaluation was intended to capture the impressions and difficulties "
    "experienced by a new user encountering the platform for the first time, "
    "which is the most ecologically valid context for assessing usability in "
    "the platform's early deployment phase."
)

add_heading(doc, "5.6.2  Session Flow", level=2)

add_body(doc,
    "Upon accessing the survey URL, participants encountered the consent and "
    "information stage. Once consent was confirmed, they selected their role "
    "and affiliation, then entered the guided tour stage. The live EDT platform "
    "was accessible throughout the session — participants could interact with "
    "it through the embedded preview or in a separate browser tab. After the "
    "tour, participants completed the three structured tasks sequentially, "
    "recording their outcomes and optional notes in the survey interface without "
    "leaving the platform context. The feature rating stage followed, "
    "succeeded by the open feedback stage. On submitting their feedback, "
    "participants were presented with the AI-generated summary and a thank-you "
    "page providing access to a link for the full platform and researcher "
    "contact information."
)

add_heading(doc, "5.6.3  Survey Progress Persistence", level=2)

add_body(doc,
    "To accommodate participants who could not complete the evaluation in a "
    "single sitting, the survey instrument implemented client-side progress "
    "persistence via the browser's localStorage API. Incomplete survey state "
    "was automatically saved at each step change, and returning participants "
    "were offered the option to resume from their last completed stage or "
    "restart from the beginning. This mechanism ensured that partial "
    "responses were not lost due to interruptions, improving data completeness "
    "without requiring participant accounts or server-side session management."
)

# ── 5.7  Feature Evaluation Framework ─────────────────────────────────────────

add_heading(doc, "5.7  Feature Evaluation Framework", level=1)

add_body(doc,
    "The selection and grouping of features for the rating stage was grounded "
    "in the platform's functional architecture. The eight rated modules "
    "collectively span the platform's primary analytical and operational "
    "capabilities: spatial understanding (3D Viewer), real-time operational "
    "data (Energy Monitoring, IAQ, Fault Detection), energy generation and "
    "storage (Solar & Battery), planning and forecasting (Scenarios, Forecast), "
    "and personalisation (Role Dashboard). Table 5.4 maps each feature to its "
    "functional category and the key user story it addresses."
)

add_table(doc,
    headers=["Feature Module", "Category", "Key User Story"],
    rows=[
        ["3D Building Viewer",
         "Spatial Understanding",
         "Navigate the building model to understand spatial layout and "
         "locate specific areas or systems."],
        ["Energy Monitoring",
         "Operational Data",
         "Inspect circuit-level power demand and identify consumption "
         "patterns in real time or through replay."],
        ["Solar & Battery",
         "Energy Generation",
         "Track solar panel output and battery storage state to understand "
         "renewable energy contribution."],
        ["Room Air Quality (IAQ)",
         "Operational Data",
         "Review temperature, humidity, and CO₂ levels by room to assess "
         "occupant comfort and wellbeing."],
        ["Fault Detection",
         "Operational Data",
         "Receive automatic alerts for equipment faults with likely causes "
         "and recommended corrective actions."],
        ["Scenarios",
         "Planning & Decision Support",
         "Model the energy impact of proposed changes (e.g., equipment "
         "replacement, scheduling adjustments)."],
        ["Forecast",
         "Planning & Decision Support",
         "View seven-day predictions for energy consumption and cost to "
         "support operational planning."],
        ["Role-Specific Dashboard",
         "Personalisation",
         "Access a curated view of the metrics most relevant to the "
         "user's declared role and responsibilities."],
    ],
    col_widths=[4.0, 3.5, 8.0]
)
add_caption(doc, "Table 5.4. Feature modules, functional categories, and key user stories.")

add_body(doc,
    "Each feature was rated on a five-point scale, with the scale endpoints "
    "labelled as 1 = not useful and 5 = very useful. This formulation "
    "intentionally captured perceived utility rather than perceived ease of "
    "use, reflecting the evaluation's emphasis on whether the platform's "
    "analytical outputs supported informed decision-making, not solely whether "
    "the interface was navigable. Ease-of-use evidence was captured separately "
    "through the task completion outcomes in Stage 4."
)

# ── 5.8  Data Collection and Storage ──────────────────────────────────────────

add_heading(doc, "5.8  Data Collection and Storage Infrastructure", level=1)

add_body(doc,
    "Survey responses were collected and persisted via a serverless backend "
    "architecture. On submission of each stage's data, a JavaScript fetch "
    "call from the client transmitted a structured JSON payload to a Netlify "
    "Functions endpoint (/.netlify/functions/save-response). The serverless "
    "function authenticated to the Supabase database using a service-role key "
    "stored as a Netlify environment variable — never exposed to the client — "
    "and issued a PostgreSQL upsert against the survey_responses table, using "
    "the session's started_at timestamp as the conflict resolution key. This "
    "upsert mechanism ensured that the initial response record and the "
    "subsequent AI-summary update were stored as a single cohesive row rather "
    "than duplicate entries."
)
add_body(doc,
    "The database schema comprised a single survey_responses table with "
    "columns capturing all quantitative and qualitative response fields, "
    "together with temporal metadata (started_at, created_at) enabling the "
    "calculation of session duration. Row-level security was enabled on the "
    "table with a restrictive policy preventing direct public access; all "
    "read and write operations were mediated exclusively by the authenticated "
    "serverless functions. The admin dashboard accessed response data through "
    "a separate Netlify Functions endpoint (/.netlify/functions/list-responses), "
    "ensuring that the service-role key was never accessible from the browser."
)
add_body(doc,
    "An interaction evidence score was computed algorithmically for each "
    "response in the admin dashboard, based on seven binary signals: whether "
    "all three tasks were answered; whether at least two tasks were completed "
    "easily; whether all eight feature ratings were provided; whether at "
    "least one substantial written comment was left; whether the AI "
    "decision-support question was answered; whether an overall rating was "
    "given; and whether a measurable session duration existed. The score, "
    "expressed as a percentage of signals achieved, served as an indicator "
    "of response quality and depth of platform engagement for use during "
    "data analysis."
)

# ── 5.9  Ethical Considerations ────────────────────────────────────────────────

add_heading(doc, "5.9  Ethical Considerations", level=1)

add_body(doc,
    "The study was designed in accordance with the ethical requirements of "
    "the University of Twente and the principles of research integrity. "
    "Participation was entirely voluntary: the consent stage clearly "
    "communicated that participants could withdraw at any stage without "
    "consequence, and that the data from any withdrawn participant would not "
    "be used in analysis. Responses were anonymised by default — no "
    "participant was required to provide their name, email address, or any "
    "other personally identifying information. The optional name and email "
    "fields in Stage 6 were explicitly labelled as optional and were included "
    "solely to facilitate researcher follow-up if the participant chose to "
    "initiate contact."
)
add_body(doc,
    "Data was stored on secure cloud infrastructure (Supabase) with "
    "access restricted to the researcher and academic supervisors via "
    "role-based key authentication. No data was shared with third parties. "
    "The AI summary generation involved transmission of response data to the "
    "OpenAI API; this was disclosed implicitly through the survey's "
    "explanation of the AI summary feature, and only anonymised, "
    "non-identifying response content was transmitted. No sensitive personal "
    "data — as defined under the General Data Protection Regulation — was "
    "collected or processed."
)
add_body(doc,
    "Participants were 18 years of age or older by confirmation at the "
    "consent stage. The study posed no known risks to participants, and "
    "participation offered no direct personal benefit beyond the intrinsic "
    "satisfaction of contributing to research aimed at improving smart "
    "building management tools. The participant information sheet and consent "
    "form were reviewed and approved as part of the thesis supervision process "
    "at the University of Twente."
)

# ── 5.10  Chapter Summary ──────────────────────────────────────────────────────

add_heading(doc, "5.10  Chapter Summary", level=1)

add_body(doc,
    "This chapter has described the design, rationale, and execution of a "
    "structured usability evaluation of the Energy Digital Twin platform "
    "developed for the GATE Institute. The evaluation employed a mixed-methods, "
    "task-based approach delivered through a purpose-built browser survey "
    "instrument. Seven survey stages collected informed consent, participant "
    "stratification data, task performance evidence, quantitative feature "
    "ratings, and qualitative open feedback from a purposively recruited "
    "sample of building stakeholders spanning six professional archetypes."
)
add_body(doc,
    "The evaluation was designed to generate both summative validation evidence "
    "— demonstrating whether the platform meets its core usability requirements "
    "— and formative diagnostic evidence to inform iterative improvement. An "
    "AI-assisted summarisation component provided immediate participant-level "
    "recaps stored alongside the full response record. A secure, serverless "
    "data infrastructure ensured that all response data was persistently and "
    "reliably stored throughout the data collection period."
)
add_body(doc,
    "The findings of the evaluation — including quantitative feature rating "
    "distributions, task completion rates, adoption intent, and thematic "
    "analysis of qualitative feedback — are presented and interpreted in "
    "Chapter 6."
)

# ── References (stub) ──────────────────────────────────────────────────────────

add_spacer(doc, 2)
add_heading(doc, "References", level=1)

refs = [
    "International Organization for Standardization. (2018). ISO 9241-11:2018 "
    "Ergonomics of human-system interaction — Part 11: Usability: Definitions "
    "and concepts. ISO.",

    "Nielsen, J. (1994). Usability Engineering. Morgan Kaufmann.",

    "Rubin, J., & Chisnell, D. (2008). Handbook of Usability Testing: How to "
    "Plan, Design, and Conduct Effective Tests (2nd ed.). Wiley.",

    "Vermesan, O., & Friess, P. (Eds.). (2014). Internet of Things — From "
    "Research and Innovation to Market Deployment. River Publishers.",

    "Pérez-Lombard, L., Ortiz, J., & Pout, C. (2008). A review on buildings "
    "energy consumption information. Energy and Buildings, 40(3), 394–398.",
]

for ref in refs:
    p = doc.add_paragraph()
    para_fmt(p, first_indent=-1.25, space_before=2, space_after=4,
             line_spacing=18)
    p.paragraph_format.left_indent = Cm(1.25)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(ref)
    set_font(run, size=11)

# ── save ───────────────────────────────────────────────────────────────────────

out_path = r"c:\Users\joanw\OneDrive - University of Twente\Desktop\My Applications\RateMyApp\Chapter_5_Usability_Evaluation.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
