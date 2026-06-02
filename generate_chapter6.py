"""
Generate Chapter 6 — Results — for the Energy Digital Twin usability evaluation.
Appends to Chapter_5_Usability_Evaluation.docx if it exists, otherwise creates fresh.
Output: Chapter_6_Results.docx
Run:    python generate_chapter6.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── helpers (same set as chapter 5 generator) ──────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False,
             italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para_fmt(para, first_indent=None, space_before=None,
             space_after=None, line_spacing=None, alignment=None):
    pf = para.paragraph_format
    if first_indent  is not None: pf.first_line_indent = Cm(first_indent)
    if space_before  is not None: pf.space_before       = Pt(space_before)
    if space_after   is not None: pf.space_after        = Pt(space_after)
    if line_spacing  is not None: pf.line_spacing        = Pt(line_spacing)
    if alignment     is not None: para.alignment         = alignment

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    if level == 1:
        pf.space_before = Pt(24); pf.space_after = Pt(12)
        run = p.add_run(text); set_font(run, size=16, bold=True)
    elif level == 2:
        pf.space_before = Pt(18); pf.space_after = Pt(8)
        run = p.add_run(text); set_font(run, size=14, bold=True)
    elif level == 3:
        pf.space_before = Pt(12); pf.space_after = Pt(6)
        run = p.add_run(text); set_font(run, size=12, bold=True, italic=True)
    return p

def add_body(doc, text, indent=True):
    p = doc.add_paragraph()
    para_fmt(p, first_indent=1.25 if indent else 0,
             space_before=0, space_after=6, line_spacing=22)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text); set_font(run)
    return p

def add_body_noindent(doc, text):
    return add_body(doc, text, indent=False)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    para_fmt(p, space_before=2, space_after=2, line_spacing=18)
    run = p.add_run(text); set_font(run)
    return p

def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    para_fmt(p, space_before=2, space_after=2, line_spacing=18)
    run = p.add_run(text); set_font(run)
    return p

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def shade_row(row, hex_color):
    for cell in row.cells:
        shade_cell(cell, hex_color)

def cell_text(cell, text, size=10.5, bold=False, italic=False,
              align=WD_ALIGN_PARAGRAPH.LEFT, color=None):
    p = cell.paragraphs[0]
    p.alignment = align
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)

def add_table(doc, headers, rows, col_widths=None,
              header_color="2563EB", alt_color="F0F4FF"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    shade_row(hdr_row, header_color)
    for i, h in enumerate(headers):
        cell_text(hdr_row.cells[i], h, size=11, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER,
                  color=(255, 255, 255))
    for ri, row_data in enumerate(rows):
        row_obj = table.rows[ri + 1]
        if ri % 2 == 1:
            shade_row(row_obj, alt_color)
        for ci, val in enumerate(row_data):
            cell_text(row_obj.cells[ci], str(val))
    if col_widths:
        for row_obj in table.rows:
            for i, w in enumerate(col_widths):
                row_obj.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_fmt(p, space_before=2, space_after=10, line_spacing=16)
    run = p.add_run(text); set_font(run, size=10, italic=True)
    return p

def add_note_box(doc, text):
    """Shaded note/callout box."""
    p = doc.add_paragraph()
    para_fmt(p, first_indent=0, space_before=6, space_after=6, line_spacing=18)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Use a single-cell table as a shaded box
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    shade_row(tbl.rows[0], "FFF7ED")
    c = tbl.rows[0].cells[0]
    c.width = Cm(15.5)
    cp = c.paragraphs[0]
    para_fmt(cp, first_indent=0, space_before=4, space_after=4, line_spacing=18)
    run = cp.add_run("📋  " + text)
    set_font(run, size=10.5, italic=True, color=(120, 53, 15))
    doc.add_paragraph()

def add_spacer(doc, lines=1):
    for _ in range(lines):
        p = doc.add_paragraph()
        para_fmt(p, space_before=0, space_after=0, line_spacing=6)

def add_placeholder_row_note(doc, note="Values to be populated from the admin dashboard after data collection."):
    p = doc.add_paragraph()
    para_fmt(p, first_indent=0, space_before=0, space_after=8, line_spacing=16)
    run = p.add_run(f"* {note}")
    set_font(run, size=9.5, italic=True, color=(100, 116, 139))


# ── document ───────────────────────────────────────────────────────────────────

doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.5)
    section.right_margin  = Cm(2.5)

# ── CHAPTER TITLE ──────────────────────────────────────────────────────────────

add_spacer(doc, 2)
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_fmt(tp, space_before=0, space_after=4)
set_font(tp.add_run("CHAPTER 6"), size=13, bold=True, color=(37, 99, 235))

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_fmt(sp, space_before=4, space_after=20)
set_font(sp.add_run(
    "Results: Usability Evaluation of the Energy Digital Twin\n"
    "for Smart Building Management"
), size=20, bold=True)

dp = doc.add_paragraph()
dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
para_fmt(dp, space_before=0, space_after=24)
set_font(dp.add_run("─" * 60), size=10, color=(148, 163, 184))

add_note_box(doc,
    "Reading this chapter: Tables marked with an asterisk (*) contain "
    "representative placeholder values derived from the survey instrument "
    "design. The researcher should replace these with the values displayed "
    "in the admin dashboard once real data collection is complete. "
    "All metric definitions and calculation methods are described in Chapter 5."
)

# ── 6.1  Introduction ──────────────────────────────────────────────────────────

add_heading(doc, "6.1  Introduction", level=1)

add_body(doc,
    "This chapter presents the results of the usability evaluation of the "
    "Energy Digital Twin (EDT) platform conducted with participants drawn "
    "from the University of Twente and GATE Institute communities. The "
    "evaluation was administered through the RateMyApp survey instrument "
    "described in Chapter 5. Results are reported in six thematic areas: "
    "participant characteristics and response quality (Section 6.2); "
    "task-based performance (Section 6.3); feature utility ratings "
    "(Section 6.4); overall satisfaction and adoption intent (Section 6.5); "
    "qualitative feedback and AI-generated response summaries (Section 6.6); "
    "and a cross-role comparative analysis (Section 6.7). A consolidated "
    "research findings snapshot is provided in Section 6.8, and the chapter "
    "concludes with a summary in Section 6.9."
)
add_body(doc,
    "All quantitative results were derived from the structured response "
    "records stored in the Supabase database and rendered by the admin "
    "dashboard. The dashboard computes aggregate statistics algorithmically "
    "from the stored JSON and relational fields; the values reported here "
    "reflect those computations directly. Qualitative results draw on the "
    "open-text fields (what_worked, what_needed, other_comments) and on the "
    "AI-generated per-participant summaries produced by the GPT-4.1-mini "
    "model during the survey's thank-you stage."
)

# ── 6.2  Participant Overview ──────────────────────────────────────────────────

add_heading(doc, "6.2  Participant Overview", level=1)

add_heading(doc, "6.2.1  Response Volume and Session Duration", level=2)

add_body(doc,
    "The evaluation collected a total of N survey responses over the data "
    "collection period. Of these, N_summary responses included an "
    "AI-generated summary, confirming that participants reached and submitted "
    "the final stage of the survey. The admin dashboard's summary statistic "
    "card reports this count directly as 'N with AI summary'. Table 6.1 "
    "presents the headline participation indicators as reported by the "
    "dashboard."
)

add_table(doc,
    headers=["Metric", "Value", "Source in Admin Dashboard"],
    rows=[
        ["Total responses",                    "N",          "Summary stats — Total Responses card"],
        ["Responses with AI summary",          "N_summary",  "Summary stats — 'N with AI summary' sub-label"],
        ["Average session duration (minutes)", "X min",      "Thesis-ready evidence — Measured sessions card"],
        ["Mean tasks answered per respondent", "X.X / 3",    "Thesis-ready evidence — Mean tasks answered"],
        ["Average interaction evidence score", "XX%",        "Interaction evidence stats — Interaction Evidence Score card"],
    ],
    col_widths=[6.0, 2.5, 7.0]
)
add_caption(doc, "Table 6.1. Headline participation indicators.*")
add_placeholder_row_note(doc)

add_body(doc,
    "Session duration was derived from the difference between the "
    "started_at timestamp — recorded when the participant first loaded the "
    "survey — and the created_at timestamp recorded on database insertion. "
    "Only sessions within the range of one to 180 minutes were considered "
    "valid; sessions outside this window were excluded from the average as "
    "they likely represent interrupted or browser-suspended sessions. The "
    "mean session duration provides contextual evidence that participants "
    "spent substantive time on the platform before submitting feedback."
)

add_heading(doc, "6.2.2  Stakeholder Role Distribution", level=2)

add_body(doc,
    "Participants self-identified their professional role from six predefined "
    "stakeholder archetypes. Table 6.2 presents the distribution of responses "
    "by role, as shown in the admin dashboard's 'Participant Roles' bar chart "
    "section. The role distribution reflects the purposive sampling strategy "
    "described in Chapter 5 and indicates the breadth of the stakeholder "
    "perspective captured by the evaluation."
)

add_table(doc,
    headers=["Stakeholder Role", "Responses (n)", "% of Total"],
    rows=[
        ["Building Director / Management",   "n₁", "x.x%"],
        ["Facilities Manager",               "n₂", "x.x%"],
        ["Sustainability Officer",           "n₃", "x.x%"],
        ["IT Staff",                         "n₄", "x.x%"],
        ["Office Worker",                    "n₅", "x.x%"],
        ["External Visitor / Stakeholder",   "n₆", "x.x%"],
        ["Total",                            "N",  "100%"],
    ],
    col_widths=[7.5, 3.0, 3.0]
)
add_caption(doc, "Table 6.2. Distribution of responses by stakeholder role.*")
add_placeholder_row_note(doc)

add_body(doc,
    "The admin dashboard additionally reports the number of distinct roles "
    "represented among all responses ('Roles Represented' card in the summary "
    "stats grid). A higher value — approaching the maximum of six — indicates "
    "that the evaluation captured a multi-perspective stakeholder sample, "
    "strengthening the ecological validity of the findings."
)

add_heading(doc, "6.2.3  Organisational Affiliation", level=2)

add_body(doc,
    "Participants were drawn from three affiliation categories: University of "
    "Twente, GATE Institute, and Other. The affiliation distribution is stored "
    "in the affiliation column of the response record and is available in the "
    "raw data table of the admin dashboard. Table 6.3 summarises the "
    "affiliation breakdown."
)

add_table(doc,
    headers=["Affiliation", "Responses (n)", "% of Total"],
    rows=[
        ["University of Twente",   "n₁", "x.x%"],
        ["GATE Institute",         "n₂", "x.x%"],
        ["Other",                  "n₃", "x.x%"],
        ["Total",                  "N",  "100%"],
    ],
    col_widths=[6.0, 3.5, 3.5]
)
add_caption(doc, "Table 6.3. Distribution of responses by organisational affiliation.*")
add_placeholder_row_note(doc)

add_heading(doc, "6.2.4  Interaction Evidence Quality", level=2)

add_body(doc,
    "The admin dashboard computes an interaction evidence score for each "
    "response based on seven binary signals: whether all three tasks were "
    "answered; whether at least two tasks were completed without reporting "
    "failure; whether all eight feature ratings were provided; whether at "
    "least one substantive open-text comment was left; whether the AI "
    "decision-support question in Task 3 was answered; whether an overall "
    "satisfaction rating was given; and whether a measurable session duration "
    "could be derived from timestamps. The score, expressed as a percentage "
    "of signals achieved, characterises the depth of platform engagement "
    "demonstrated by each response."
)
add_body(doc,
    "Table 6.4 presents the distribution of responses across interaction "
    "score bands, as reported in the 'Interaction score spread' bar chart "
    "of the admin dashboard. A concentration of responses in the upper "
    "bands (60–100%) indicates that the majority of participants engaged "
    "substantively with the platform rather than providing superficial "
    "responses."
)

add_table(doc,
    headers=["Interaction Score Band", "Responses (n)", "% of Total", "Interpretation"],
    rows=[
        ["0–39%",    "n₁", "x.x%", "Minimal engagement — few signals achieved"],
        ["40–59%",   "n₂", "x.x%", "Moderate engagement — partial task/rating completion"],
        ["60–79%",   "n₃", "x.x%", "Good engagement — most signals achieved"],
        ["80–100%",  "n₄", "x.x%", "Full engagement — near-complete evidence set"],
        ["Total",    "N",  "100%",  ""],
    ],
    col_widths=[3.5, 2.5, 2.5, 7.0]
)
add_caption(doc, "Table 6.4. Distribution of responses by interaction evidence score band.*")
add_placeholder_row_note(doc)

add_body(doc,
    "Four additional thesis-ready indicators are reported in the admin "
    "dashboard under the 'Thesis-ready evidence' section. These are: the "
    "number of participants who completed all three tasks (task participation); "
    "the number who rated all eight features (deep feedback); the number who "
    "answered the AI decision-support question in Task 3 (AI support "
    "engagement); and the average session duration for measurably timed "
    "sessions (measured sessions). These indicators collectively substantiate "
    "that participants interacted with the platform in a hands-on, task-driven "
    "manner rather than responding to the survey from prior knowledge alone. "
    "Table 6.5 presents these values."
)

add_table(doc,
    headers=["Thesis-ready Indicator", "Value", "Interpretation"],
    rows=[
        ["Completed all 3 tasks",                 "n / N",   "Proportion who answered all guided tasks"],
        ["Completed at least 2 tasks",            "n / N",   "Partial completion — broad task engagement"],
        ["Rated all 8 features",                  "n / N",   "Deep feature-level feedback coverage"],
        ["Left written reflection",               "n / N",   "Qualitative evidence available for analysis"],
        ["Answered AI support question (Task 3)", "n / N",   "Reached and interpreted the final task"],
        ["Average session duration",              "X min",   "Mean timed session length"],
    ],
    col_widths=[6.5, 2.5, 6.5]
)
add_caption(doc, "Table 6.5. Thesis-ready interaction evidence indicators.*")
add_placeholder_row_note(doc)

# ── 6.3  Task-Based Performance ────────────────────────────────────────────────

add_heading(doc, "6.3  Task-Based Performance Results", level=1)

add_body(doc,
    "Task performance was assessed through three structured scenarios, each "
    "requiring participants to interact with the live EDT platform and "
    "self-report their outcome. For each task, participants chose from three "
    "completion descriptors: 'Yes, easily', 'Yes, with some difficulty', or "
    "'No, I couldn't find it'. Task 1 and Task 2 included an additional "
    "follow-up question on perceived clarity or ease, and Task 3 included a "
    "question on the perceived helpfulness of the AI assistant. Participants "
    "were also offered an optional free-text field to record qualitative "
    "observations about each task."
)
add_body(doc,
    "The admin dashboard reports task outcomes in two complementary ways: "
    "the 'Task Completion' section shows the percentage of participants who "
    "reported completing each task easily; and the 'Task outcome comparison' "
    "stacked bar chart in the Response Graphs section shows the full "
    "distribution across all three outcome categories plus non-response. "
    "The following sub-sections report both representations."
)

add_heading(doc, "6.3.1  Task 1: Explore the Building and Select a View", level=2)

add_body(doc,
    "Task 1 required participants to navigate the 3D building model, zoom, "
    "rotate, and explore different areas of the building. The follow-up "
    "question asked participants to rate how easily they could find and "
    "navigate to different areas. Table 6.6 presents the distribution of "
    "Task 1 outcomes."
)

add_table(doc,
    headers=["Outcome", "n", "% of Total", "Follow-up: Navigation Ease"],
    rows=[
        ["Completed easily",          "n₁", "xx%", "Very easy / Somewhat easy / Difficult"],
        ["Completed with difficulty", "n₂", "xx%", ""],
        ["Could not complete",        "n₃", "xx%", ""],
        ["No result recorded",        "n₄", "xx%", ""],
        ["Total",                     "N",  "100%",""],
    ],
    col_widths=[5.0, 1.5, 2.0, 7.0]
)
add_caption(doc, "Table 6.6. Task 1 outcome distribution — 3D navigation.*")
add_placeholder_row_note(doc)

add_body(doc,
    "The admin dashboard reports the 'easy rate' — the percentage of "
    "respondents who selected 'Completed easily' — as the headline Task 1 "
    "metric in the summary stats card ('Task 1 Easy Rate'). A high easy rate "
    "for Task 1 would suggest that the platform's spatial navigation is "
    "intuitive for first-time users across stakeholder groups, which is a "
    "foundational usability requirement for a building management digital twin."
)

add_heading(doc, "6.3.2  Task 2: Inspect Building Conditions", level=2)

add_body(doc,
    "Task 2 directed participants to use the circuit dropdown, replay "
    "functionality, IAQ room panels, and heatmap overlays to understand "
    "current and historical building conditions. The follow-up question "
    "asked participants to assess how clearly the platform communicated "
    "building conditions. Table 6.7 presents the Task 2 outcome distribution."
)

add_table(doc,
    headers=["Outcome", "n", "% of Total", "Follow-up: Condition Clarity"],
    rows=[
        ["Completed easily",          "n₁", "xx%", "Very clear / Somewhat clear / Unclear"],
        ["Completed with difficulty", "n₂", "xx%", ""],
        ["Could not complete",        "n₃", "xx%", ""],
        ["No result recorded",        "n₄", "xx%", ""],
        ["Total",                     "N",  "100%",""],
    ],
    col_widths=[5.0, 1.5, 2.0, 7.0]
)
add_caption(doc, "Table 6.7. Task 2 outcome distribution — dashboard and heatmap inspection.*")
add_placeholder_row_note(doc)

add_heading(doc, "6.3.3  Task 3: Use the AI Assistant for Decision Support", level=2)

add_body(doc,
    "Task 3 asked participants to interact with the conversational AI "
    "assistant embedded in the platform, including spatial queries, data "
    "queries, and questions about recommended actions. The follow-up "
    "question assessed whether the AI's responses helped participants "
    "understand what actions might be needed. Table 6.8 presents the Task 3 "
    "outcome distribution, including responses to the AI helpfulness question."
)

add_table(doc,
    headers=["Outcome", "n", "% of Total", "AI Helpfulness Rating"],
    rows=[
        ["Completed easily",          "n₁", "xx%", "Very helpful / Somewhat helpful / Not helpful"],
        ["Completed with difficulty", "n₂", "xx%", ""],
        ["Could not complete",        "n₃", "xx%", ""],
        ["No result recorded",        "n₄", "xx%", ""],
        ["Total",                     "N",  "100%",""],
    ],
    col_widths=[5.0, 1.5, 2.0, 7.0]
)
add_caption(doc, "Table 6.8. Task 3 outcome distribution — AI assistant for decision support.*")
add_placeholder_row_note(doc)

add_heading(doc, "6.3.4  Cross-Task Comparison", level=2)

add_body(doc,
    "Table 6.9 provides a side-by-side comparison of the easy-completion "
    "rates and scenario-judgement counts across all three tasks, as rendered "
    "in the 'Task Completion' section and the stacked task outcome chart of "
    "the admin dashboard. This comparison allows the relative difficulty of "
    "each task scenario to be assessed and surfaces which platform capabilities "
    "present the greatest navigability or comprehension challenges."
)

add_table(doc,
    headers=["Task", "Easy Rate (%)", "Difficult (%)", "Not Completed (%)", "Interpretation"],
    rows=[
        ["Task 1 — 3D Navigation",          "xx%", "xx%", "xx%",
         "Spatial orientation of the 3D model"],
        ["Task 2 — Dashboard Inspection",   "xx%", "xx%", "xx%",
         "Data clarity and replay usability"],
        ["Task 3 — AI Decision Support",    "xx%", "xx%", "xx%",
         "Perceived utility of AI assistant"],
    ],
    col_widths=[4.5, 2.5, 2.5, 2.5, 3.5]
)
add_caption(doc, "Table 6.9. Cross-task completion rate comparison.*")
add_placeholder_row_note(doc)

# ── 6.4  Feature Utility Ratings ───────────────────────────────────────────────

add_heading(doc, "6.4  Feature Utility Ratings", level=1)

add_heading(doc, "6.4.1  Overview of Feature Ratings", level=2)

add_body(doc,
    "Following the task scenarios, participants rated each of the eight "
    "feature modules on a five-point Likert scale (1 = not useful, "
    "5 = very useful). The admin dashboard's 'Feature Ratings' section "
    "displays these averages in descending order of mean score, enabling "
    "immediate identification of the highest- and lowest-rated features. "
    "Table 6.10 presents the mean rating and number of raters for each "
    "feature module, ordered by descending mean score as the admin dashboard "
    "renders them."
)

add_table(doc,
    headers=["Rank", "Feature Module", "Mean Rating (1–5)", "n Raters",
             "Selected as Most Useful", "Selected as Needs Work"],
    rows=[
        ["1", "3D Building Viewer",        "x.xx", "n", "n", "n"],
        ["2", "Energy Monitoring",         "x.xx", "n", "n", "n"],
        ["3", "Fault Detection",           "x.xx", "n", "n", "n"],
        ["4", "Room Air Quality (IAQ)",    "x.xx", "n", "n", "n"],
        ["5", "Role-Specific Dashboard",   "x.xx", "n", "n", "n"],
        ["6", "Solar & Battery",           "x.xx", "n", "n", "n"],
        ["7", "Scenarios",                 "x.xx", "n", "n", "n"],
        ["8", "Forecast",                  "x.xx", "n", "n", "n"],
    ],
    col_widths=[1.2, 4.5, 3.0, 1.5, 2.8, 2.5]
)
add_caption(doc,
    "Table 6.10. Feature utility ratings sorted by mean score (descending). "
    "Rank order will reflect actual admin dashboard ordering.*")
add_placeholder_row_note(doc)

add_body(doc,
    "The 'Most Useful' and 'Needs Work' columns reflect the counts from "
    "the forced-choice dropdown questions at the end of the feature rating "
    "stage, as displayed in the admin dashboard's 'Stakeholder interaction "
    "by feature' section. These columns complement the Likert mean by "
    "capturing deliberate comparative preference: a feature may receive a "
    "moderate mean rating but still attract a high 'most useful' count if "
    "it is perceived as distinctively more valuable than alternatives."
)

add_heading(doc, "6.4.2  Highest-Rated Features", level=2)

add_body(doc,
    "The admin dashboard's 'Research findings snapshot' section automatically "
    "identifies the strongest perceived feature as the one with the highest "
    "'Most Useful' selection count, using mean rating as a tiebreaker. "
    "[Feature name] was identified as the strongest perceived feature, with "
    "a mean rating of X.XX/5 and selected as most useful by N participant(s). "
    "This finding suggests that [feature description and interpretation — to "
    "be written by researcher based on actual values]."
)

add_heading(doc, "6.4.3  Features Requiring Improvement", level=2)

add_body(doc,
    "The feature with the highest improvement demand — identified by the "
    "greatest 'Needs Work' selection count, with mean rating as a "
    "tiebreaker — was [Feature name], flagged by N participant(s) and "
    "carrying a mean rating of X.XX/5. Qualitative comments associated "
    "with this feature, available in the optional per-feature comment "
    "fields of the raw response data, provided additional insight into "
    "the specific nature of the improvement demand. [Summary of relevant "
    "qualitative comments — to be written by researcher]."
)

add_heading(doc, "6.4.4  Feature Interest Relative to Improvement Demand", level=2)

add_body(doc,
    "Figure 6.1 (replicated here as Table 6.11) shows the relationship "
    "between the 'Most Useful' and 'Needs Work' selection counts for each "
    "feature module, as displayed in the 'Stakeholder interaction by feature' "
    "section of the admin dashboard. Features appearing high on both "
    "dimensions represent capabilities that are perceived as valuable but "
    "not yet fully realised — they are strategically important targets for "
    "the next development iteration."
)

add_table(doc,
    headers=["Feature Module", "Most Useful Count", "Needs Work Count",
             "Strategic Classification"],
    rows=[
        ["3D Building Viewer",      "n", "n", "High value / [low|high] demand"],
        ["Energy Monitoring",       "n", "n", "High value / [low|high] demand"],
        ["Fault Detection",         "n", "n", "High value / [low|high] demand"],
        ["Room Air Quality (IAQ)",  "n", "n", "High value / [low|high] demand"],
        ["Role-Specific Dashboard", "n", "n", "High value / [low|high] demand"],
        ["Solar & Battery",         "n", "n", "[Low|High] value / [low|high] demand"],
        ["Scenarios",               "n", "n", "[Low|High] value / [low|high] demand"],
        ["Forecast",                "n", "n", "[Low|High] value / [low|high] demand"],
    ],
    col_widths=[4.5, 2.8, 2.8, 5.4]
)
add_caption(doc,
    "Table 6.11. Feature interest and improvement demand (replication of admin dashboard "
    "'Stakeholder interaction by feature' chart).*")
add_placeholder_row_note(doc)

# ── 6.5  Overall Satisfaction and Adoption Intent ─────────────────────────────

add_heading(doc, "6.5  Overall Platform Satisfaction and Adoption Intent", level=1)

add_heading(doc, "6.5.1  Overall Satisfaction Rating", level=2)

add_body(doc,
    "Participants rated their overall impression of the platform on a "
    "five-point emoji-anchored scale, with labels ranging from 'Needs a "
    "lot of work' (1) to 'Exceptional' (5). The admin dashboard reports "
    "the mean overall rating in the summary stats grid ('Avg Overall "
    "Rating' card) and displays the full distribution in the 'Overall "
    "ratings pie' chart and 'Overall rating distribution' bar chart. "
    "Table 6.12 presents this distribution."
)

add_table(doc,
    headers=["Rating", "Label", "n", "% of Total"],
    rows=[
        ["1", "Needs a lot of work",   "n₁", "xx%"],
        ["2", "Below expectations",    "n₂", "xx%"],
        ["3", "Meets expectations",    "n₃", "xx%"],
        ["4", "Exceeds expectations",  "n₄", "xx%"],
        ["5", "Exceptional",           "n₅", "xx%"],
        ["",  "Mean (± SD)",           "x.xx", ""],
    ],
    col_widths=[2.0, 4.5, 2.0, 2.5]
)
add_caption(doc, "Table 6.12. Overall platform satisfaction rating distribution.*")
add_placeholder_row_note(doc)

add_body(doc,
    "The mean overall rating of X.XX/5 [interpretation: e.g., positions the "
    "platform at a level that meets or exceeds user expectations for a "
    "research prototype at this stage of development]. The distribution shape "
    "— as visible in the admin dashboard's pie chart — indicates [whether "
    "ratings clustered at higher values, were normally distributed, or showed "
    "bimodal patterns reflecting divergent stakeholder perspectives]."
)

add_heading(doc, "6.5.2  Use Intention", level=2)

add_body(doc,
    "Participants' willingness to incorporate the EDT platform into their "
    "daily work was assessed through a four-option forced-choice question. "
    "The admin dashboard reports the results in both the 'Adoption signal' "
    "card of the Research Findings Snapshot and the 'Daily-use intention' "
    "bar chart and 'Use-intention pie' chart in the Response Graphs section. "
    "Table 6.13 presents the adoption intent distribution."
)

add_table(doc,
    headers=["Response", "n", "% of Total", "Classification"],
    rows=[
        ["Yes, regularly",   "n₁", "xx%", "Positive — strong adoption intent"],
        ["Yes, occasionally","n₂", "xx%", "Positive — moderate adoption intent"],
        ["Probably not",     "n₃", "xx%", "Negative — conditional rejection"],
        ["No",               "n₄", "xx%", "Negative — outright rejection"],
        ["Total positive",   "n₁+n₂", "xx%", ""],
        ["Total negative",   "n₃+n₄", "xx%", ""],
    ],
    col_widths=[4.0, 1.5, 2.5, 7.5]
)
add_caption(doc, "Table 6.13. Use-intention distribution.*")
add_placeholder_row_note(doc)

add_body(doc,
    "The admin dashboard's 'Adoption signal' card summarises this as a "
    "combined positive count: 'N/Total said they would use the system at "
    "least occasionally.' [Researcher interpretation of what this adoption "
    "signal implies for the platform's readiness for operational deployment "
    "or continued iterative development — to be written after data collection]."
)

# ── 6.6  Qualitative Feedback ──────────────────────────────────────────────────

add_heading(doc, "6.6  Qualitative Feedback and AI-Generated Summaries", level=1)

add_heading(doc, "6.6.1  What Participants Found Most Useful", level=2)

add_body(doc,
    "Participants were invited to describe what they found most useful or "
    "impressive about the platform in an open-text field (what_worked). "
    "This field is visible in the raw data table and within the individual "
    "participant insight cards of the admin dashboard. The following themes "
    "emerged from a first-pass review of these responses:"
)
add_bullet(doc, "[Theme 1 — to be populated from admin dashboard qualitative data]")
add_bullet(doc, "[Theme 2 — to be populated from admin dashboard qualitative data]")
add_bullet(doc, "[Theme 3 — to be populated from admin dashboard qualitative data]")
add_body_noindent(doc,
    "Representative quotations from the what_worked field are provided below. "
    "These are drawn directly from the participant insight cards in the admin "
    "dashboard, which display each participant's verbatim text alongside their "
    "role, overall rating, and interaction score."
)

# blockquote style for representative quotes
for placeholder in [
    '"[Verbatim participant quote from what_worked field — '
    'copy from admin dashboard participant insight card #N, Role: X]"',

    '"[Verbatim participant quote from what_worked field — '
    'copy from admin dashboard participant insight card #N, Role: X]"',
]:
    p = doc.add_paragraph()
    para_fmt(p, first_indent=0, space_before=4, space_after=4, line_spacing=20)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent  = Cm(2.0)
    p.paragraph_format.right_indent = Cm(1.5)
    run = p.add_run(placeholder)
    set_font(run, size=11, italic=True, color=(71, 85, 105))

add_heading(doc, "6.6.2  Main Friction Points and Improvement Requests", level=2)

add_body(doc,
    "The what_needed field and the 'Main Friction' block in each participant "
    "insight card captured what participants found confusing, missing, or "
    "frustrating. These responses represent formative usability evidence of "
    "direct relevance to the next design iteration. The following themes "
    "were identified:"
)
add_bullet(doc, "[Friction theme 1 — to be populated from admin dashboard]")
add_bullet(doc, "[Friction theme 2 — to be populated from admin dashboard]")
add_bullet(doc, "[Friction theme 3 — to be populated from admin dashboard]")

add_body_noindent(doc,
    "The admin dashboard's 'Needs Work' block in each participant card "
    "provides the verbatim text, while the 'Highest improvement demand' "
    "card in the Research Findings Snapshot identifies the single feature "
    "most frequently nominated for improvement. These two views together "
    "provide both specific textual evidence and a quantitative priority "
    "signal for the development team."
)

add_heading(doc, "6.6.3  AI-Generated Participant Summaries", level=2)

add_body(doc,
    "At the conclusion of the survey, each participant's responses were "
    "processed by the GPT-4.1-mini language model to generate a personalised "
    "feedback summary addressed to the development team. The model was "
    "prompted to identify what the participant valued most, the main pain "
    "point, and one specific actionable recommendation, expressed in plain "
    "English without bullet points. These summaries were stored in the "
    "ai_summary column of the database and are displayed both in the "
    "dedicated 'AI-Generated Summaries' section and within each participant's "
    "insight card in the admin dashboard."
)
add_body(doc,
    "The following representative AI-generated summaries illustrate the "
    "range of feedback themes captured across stakeholder groups. Each "
    "summary is presented with the participant's declared role and response "
    "number as displayed in the admin dashboard."
)

# AI summary cards
for placeholder, role in [
    ("[AI summary text for Response #N — copy verbatim from admin dashboard "
     "AI-Generated Summaries section or participant insight card]",
     "Role: [e.g., Facilities Manager] — Response #N"),

    ("[AI summary text for Response #N — copy verbatim from admin dashboard "
     "AI-Generated Summaries section or participant insight card]",
     "Role: [e.g., Sustainability Officer] — Response #N"),

    ("[AI summary text for Response #N — copy verbatim from admin dashboard "
     "AI-Generated Summaries section or participant insight card]",
     "Role: [e.g., Building Director] — Response #N"),
]:
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    shade_row(tbl.rows[0], "F0FDF4")
    c = tbl.rows[0].cells[0]
    c.width = Cm(15.5)
    label_p = c.paragraphs[0]
    para_fmt(label_p, first_indent=0, space_before=4, space_after=2, line_spacing=16)
    label_run = label_p.add_run(role)
    set_font(label_run, size=10, bold=True, color=(22, 163, 74))
    body_p = c.add_paragraph()
    para_fmt(body_p, first_indent=0, space_before=0, space_after=4, line_spacing=18)
    body_run = body_p.add_run(placeholder)
    set_font(body_run, size=11, italic=True, color=(30, 41, 59))
    doc.add_paragraph()

add_body(doc,
    "The AI summaries provide a rapid-scanning layer that complements "
    "statistical analysis: while aggregate ratings identify what was valued "
    "or problematic at a feature level, the summaries surface the contextual "
    "reasoning and specific experiential detail that statistical tables cannot "
    "capture. The researcher used these summaries as a first-pass thematic "
    "orientation before conducting deeper qualitative analysis of the verbatim "
    "text fields."
)

# ── 6.7  Stakeholder Evidence by Role ─────────────────────────────────────────

add_heading(doc, "6.7  Comparative Analysis by Stakeholder Role", level=1)

add_body(doc,
    "The admin dashboard's 'Stakeholder evidence by role' section computes "
    "four comparative metrics for each stakeholder archetype: the number of "
    "participants who completed all three tasks; the mean overall satisfaction "
    "rating; the number who expressed positive adoption intent; and the mean "
    "interaction evidence score. These metrics are presented in Table 6.14, "
    "enabling a role-stratified view of platform usability."
)

add_table(doc,
    headers=["Role", "n", "All Tasks (%)",
             "Avg Overall (/5)", "Positive Adoption (%)", "Avg Interaction Score (%)"],
    rows=[
        ["Director / Management",   "n", "xx%", "x.xx", "xx%", "xx%"],
        ["Facilities Manager",      "n", "xx%", "x.xx", "xx%", "xx%"],
        ["Sustainability Officer",  "n", "xx%", "x.xx", "xx%", "xx%"],
        ["IT Staff",                "n", "xx%", "x.xx", "xx%", "xx%"],
        ["Office Worker",           "n", "xx%", "x.xx", "xx%", "xx%"],
        ["External Visitor",        "n", "xx%", "x.xx", "xx%", "xx%"],
    ],
    col_widths=[4.2, 1.0, 2.2, 2.5, 2.8, 2.8]
)
add_caption(doc,
    "Table 6.14. Stakeholder evidence by role (source: admin dashboard "
    "'Stakeholder evidence by role' section).*")
add_placeholder_row_note(doc)

add_body(doc,
    "Role-level comparisons should be interpreted with caution where the "
    "number of responses per role is small; differences in aggregate scores "
    "for roles represented by one or two participants do not carry sufficient "
    "statistical weight to draw definitive conclusions. Nevertheless, "
    "systematic patterns — for example, consistently higher interaction scores "
    "among operationally focused roles such as Facilities Manager compared to "
    "strategically focused roles such as Building Director — can suggest that "
    "the platform's current feature set resonates more strongly with "
    "operational users and that additional design work may be needed to serve "
    "senior stakeholder information needs."
)

# ── 6.8  Research Findings Snapshot ───────────────────────────────────────────

add_heading(doc, "6.8  Research Findings Snapshot", level=1)

add_body(doc,
    "The admin dashboard's 'Research findings snapshot' section provides a "
    "condensed three-card narrative summary automatically computed from the "
    "collected data: the strongest perceived feature (by most-useful count "
    "and mean rating), the feature with the highest improvement demand (by "
    "needs-work count and mean rating), and the adoption signal (positive "
    "intent count out of total). Table 6.15 replicates these findings in "
    "tabular form for incorporation into the thesis."
)

add_table(doc,
    headers=["Finding Dimension", "Value", "Detail"],
    rows=[
        ["Strongest perceived feature",
         "[Feature name]",
         "Avg X.XX/5; selected as most useful by N participant(s)."],
        ["Highest improvement demand",
         "[Feature name]",
         "Avg X.XX/5; flagged as needing work by N participant(s)."],
        ["Positive adoption signal",
         "N / Total",
         "Would use at least occasionally: N. Regular use intent: N."],
        ["Mean overall satisfaction",
         "X.XX / 5",
         "Distribution: see Table 6.12."],
        ["Mean interaction evidence score",
         "XX%",
         "Based on 7 binary engagement signals per response."],
    ],
    col_widths=[5.0, 3.5, 7.0]
)
add_caption(doc,
    "Table 6.15. Research findings snapshot (source: admin dashboard "
    "'Research findings snapshot' and summary stats sections).*")
add_placeholder_row_note(doc)

add_body(doc,
    "These findings are discussed in the context of the evaluation objectives "
    "stated in Chapter 5 and in relation to the existing literature on smart "
    "building interface design in Chapter 7. Taken together, they indicate "
    "[high-level interpretive statement — to be written by researcher after "
    "reviewing actual results]: the platform's [area of strength] demonstrates "
    "clear value to its target stakeholders, while [area for improvement] "
    "represents the most actionable target for the next design iteration."
)

# ── 6.9  Chapter Summary ──────────────────────────────────────────────────────

add_heading(doc, "6.9  Chapter Summary", level=1)

add_body(doc,
    "This chapter has presented the quantitative and qualitative results of "
    "the usability evaluation of the Energy Digital Twin platform. Across N "
    "responses spanning six stakeholder roles, the evaluation produced evidence "
    "on three-dimensional spatial navigation (Task 1), dashboard and heatmap "
    "comprehension (Task 2), and AI-assisted decision support (Task 3); Likert "
    "utility ratings for all eight feature modules; overall satisfaction and "
    "adoption intent distributions; verbatim qualitative feedback; and "
    "AI-generated per-participant summaries."
)
add_body(doc,
    "The results were derived from structured records stored in a Supabase "
    "database and computed by the RateMyApp admin dashboard, which rendered "
    "aggregate statistics, role-stratified comparisons, visual charts, and "
    "the full set of individual participant insight cards — each augmented "
    "with its AI-generated summary. This infrastructure ensured that all "
    "reported values are directly traceable to the source data, supporting "
    "research transparency and reproducibility."
)
add_body(doc,
    "Chapter 7 interprets these findings in light of the research objectives "
    "stated in Chapter 5, discussing their implications for the design and "
    "deployment of the EDT platform and situating them within the broader "
    "literature on digital twins, smart building management, and human-computer "
    "interaction in energy-related decision support systems."
)

# ── save ───────────────────────────────────────────────────────────────────────

out_path = (
    r"c:\Users\joanw\OneDrive - University of Twente"
    r"\Desktop\My Applications\RateMyApp"
    r"\Chapter_6_Results.docx"
)
doc.save(out_path)
print(f"Saved: {out_path}")
